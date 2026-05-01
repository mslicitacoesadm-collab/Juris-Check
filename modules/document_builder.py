from __future__ import annotations
from io import BytesIO
import html
import re
from datetime import datetime
from typing import Iterable

from modules.report_builder import build_audit_summary, build_export_rows, recommended_action


def _clean(text: str) -> str:
    return re.sub(r'\s+', ' ', str(text or '')).strip()


def _paragraphs(text: str) -> list[str]:
    parts = re.split(r'\n\s*\n+|\r\n\r\n+', text or '')
    if len(parts) <= 1:
        parts = re.split(r'\n+', text or '')
    return [p.strip() for p in parts if p and p.strip()]


def _suggestion(item: dict) -> dict:
    return item.get('correcao_sugerida') or item.get('matched_record') or {}


def build_change_log(analysis: dict) -> list[dict]:
    rows = []
    for idx, item in enumerate(analysis.get('citation_results', []) or [], start=1):
        sug = _suggestion(item)
        status = item.get('status')
        if status == 'valida_compatível':
            action = 'Validado'
            after = item.get('raw', '')
        elif sug:
            action = 'Correção/Reforço sugerido'
            after = sug.get('citacao_curta', '')
        else:
            action = 'Revisão manual necessária'
            after = 'Sem substituição automática segura'
        rows.append({
            'item': idx,
            'before': item.get('raw', ''),
            'after': after,
            'line': item.get('linha', ''),
            'status': item.get('status_label', ''),
            'confidence': item.get('grau_confianca', ''),
            'action': action,
            'reason': item.get('tipo_erro', ''),
            'basis': sug.get('fundamento_curto', '') if isinstance(sug, dict) else '',
            'recommended_action': recommended_action(item),
        })
    return rows


def build_revised_text(piece_text: str, analysis: dict, mode: str = 'premium') -> str:
    """Gera texto fiel: não inventa correções; marca pontos e acrescenta log técnico."""
    text = piece_text or ''
    for item in analysis.get('citation_results', []) or []:
        raw = item.get('raw')
        if raw and item.get('status') != 'valida_compatível':
            sug = _suggestion(item)
            tag = f"[REVISAR DECIFRA: {raw}"
            if sug.get('citacao_curta'):
                tag += f" | sugestão/reforço: {sug.get('citacao_curta')}"
            tag += "]"
            text = text.replace(raw, tag, 1)
    logs = build_change_log(analysis)
    if logs:
        appendix = ['\n\n---', 'REGISTRO TÉCNICO DAS ALTERAÇÕES / VALIDAÇÕES']
        for row in logs:
            appendix.append(f"Item {row['item']}: {row['before']} -> {row['after']} | {row['action']} | {row['reason']}")
        text += '\n'.join(appendix)
    return text


def build_marked_text(piece_text: str, analysis: dict) -> str:
    return build_revised_text(piece_text, analysis, mode='marked')


def _set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text or ''))
    run.bold = bold


def _shade_cell(cell, fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def _add_small_note(paragraph, text: str):
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = None


def build_docx_bytes(text: str, analysis: dict, title: str = 'DECIFRA - Peça revisada', marked: bool = False):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(10.5)

    summary = build_audit_summary(analysis)
    changes = build_change_log(analysis)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(7, 29, 53)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Relatório gerado para apoio técnico. Revise juridicamente antes do protocolo.')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(102, 112, 133)

    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    labels = ['Tipo da peça', 'Referências', 'Pontos de revisão', 'Risco']
    vals = [summary['tipo_peca'], summary['total_referencias'], summary['pontos_revisao'], summary['risco']]
    for i, label in enumerate(labels):
        _set_cell_text(hdr[i], f'{label}\n{vals[i]}', True)
        _shade_cell(hdr[i], 'EAF2FA')
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_heading('1. Resumo executivo da auditoria', level=1)
    doc.add_paragraph(summary['justificativa_risco'])
    doc.add_paragraph(f"Referências validadas: {summary['referencias_validadas']} | Teses sugeridas: {summary['teses_sugeridas']} | Documento: {analysis.get('file_name','peça enviada')}")

    doc.add_heading('2. Registro claro do que foi verificado e corrigido', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    heads = ['Item', 'Antes', 'Depois / Sugestão', 'Status', 'Ação', 'Motivo']
    for i, h in enumerate(heads):
        _set_cell_text(table.rows[0].cells[i], h, True)
        _shade_cell(table.rows[0].cells[i], '071D35')
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    if not changes:
        cells = table.add_row().cells
        _set_cell_text(cells[0], '-')
        _set_cell_text(cells[1], 'Nenhuma citação formal detectada')
        _set_cell_text(cells[2], 'Sem alteração automática')
        _set_cell_text(cells[3], 'Indeterminado')
        _set_cell_text(cells[4], 'Usar busca por tese')
        _set_cell_text(cells[5], 'O texto não apresentou padrão formal de acórdão/súmula reconhecido')
    for row in changes:
        cells = table.add_row().cells
        values = [row['item'], row['before'], row['after'], row['status'], row['action'], row['reason']]
        for i, val in enumerate(values):
            _set_cell_text(cells[i], val)

    doc.add_heading('3. Fundamentação resumida das sugestões', level=1)
    has_basis = False
    for row in changes:
        if row.get('basis'):
            has_basis = True
            p = doc.add_paragraph()
            p.add_run(f"Item {row['item']} — {row['after']}: ").bold = True
            p.add_run(_clean(row['basis'])[:1200])
    if not has_basis:
        doc.add_paragraph('Não houve fundamento automático seguro além da validação por número/ano ou não foram localizados precedentes próximos na base disponível.')

    doc.add_heading('4. Peça com marcações de revisão', level=1)
    doc.add_paragraph('Os trechos marcados como [REVISAR DECIFRA] indicam pontos que exigem conferência ou substituição. Trechos sem marcação não foram alterados automaticamente.')
    revised = build_revised_text(text, analysis)
    for para in _paragraphs(revised)[:900]:
        p = doc.add_paragraph()
        if '[REVISAR DECIFRA:' in para:
            p.style = doc.styles['Normal']
            run = p.add_run(para)
            run.bold = True
            run.font.color.rgb = RGBColor(180, 35, 24)
        else:
            p.add_run(para)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run('DECIFRA Licitações • Auditoria técnica de citações e precedentes')
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(102, 112, 133)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_audit_docx_bytes(analysis: dict, title: str = 'DECIFRA - Relatório de auditoria'):
    return build_docx_bytes(analysis.get('piece_text', ''), analysis, title=title, marked=True)


def build_pdf_bytes(text: str, analysis: dict, title: str = 'DECIFRA - Relatório técnico'):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, rightMargin=1.5*cm, leftMargin=1.5*cm, topMargin=1.4*cm, bottomMargin=1.4*cm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='SmallMuted', parent=styles['Normal'], fontSize=8, textColor=colors.HexColor('#667085'), leading=10))
    styles.add(ParagraphStyle(name='BodyJust', parent=styles['Normal'], fontSize=9.2, leading=12))
    story = []
    summary = build_audit_summary(analysis)
    story.append(Paragraph(f'<b>{html.escape(title)}</b>', styles['Title']))
    story.append(Paragraph('Relatório técnico de apoio. Revise juridicamente antes do protocolo.', styles['SmallMuted']))
    story.append(Spacer(1, 10))
    metrics = [[
        f"Tipo<br/><b>{html.escape(str(summary['tipo_peca']))}</b>",
        f"Referências<br/><b>{summary['total_referencias']}</b>",
        f"Pontos de revisão<br/><b>{summary['pontos_revisao']}</b>",
        f"Risco<br/><b>{html.escape(summary['risco'])}</b>",
    ]]
    table = Table([[Paragraph(x, styles['Normal']) for x in metrics[0]]], colWidths=[4.1*cm, 3.2*cm, 3.5*cm, 3.2*cm])
    table.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#EAF2FA')), ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#D0D5DD')), ('PADDING', (0,0), (-1,-1), 8)]))
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(Paragraph('<b>Resumo executivo</b>', styles['Heading2']))
    story.append(Paragraph(html.escape(summary['justificativa_risco']), styles['BodyJust']))

    rows = [['Item', 'Antes', 'Depois/Sugestão', 'Status', 'Ação']]
    for row in build_change_log(analysis):
        rows.append([str(row['item']), row['before'], row['after'], row['status'], row['action']])
    if len(rows) == 1:
        rows.append(['-', 'Nenhuma referência formal detectada', 'Sem alteração automática', 'Indeterminado', 'Usar busca por tese'])
    story.append(Spacer(1, 8))
    story.append(Paragraph('<b>Registro do que foi verificado</b>', styles['Heading2']))
    trows = [[Paragraph(html.escape(str(c))[:900], styles['SmallMuted' if i else 'Normal']) for i, c in enumerate(row)] for row in rows]
    t = Table(trows, repeatRows=1, colWidths=[1.1*cm, 3.6*cm, 4.0*cm, 2.5*cm, 3.0*cm])
    t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.HexColor('#071D35')), ('TEXTCOLOR', (0,0), (-1,0), colors.white), ('GRID', (0,0), (-1,-1), 0.35, colors.HexColor('#D0D5DD')), ('VALIGN', (0,0), (-1,-1), 'TOP'), ('PADDING', (0,0), (-1,-1), 5)]))
    story.append(t)

    story.append(PageBreak())
    story.append(Paragraph('<b>Peça com marcações</b>', styles['Heading2']))
    revised = build_revised_text(text, analysis)
    for para in _paragraphs(revised[:90000])[:500]:
        safe = html.escape(para)
        if '[REVISAR DECIFRA:' in para:
            safe = f'<font color="#B42318"><b>{safe}</b></font>'
        story.append(Paragraph(safe, styles['BodyJust']))
        story.append(Spacer(1, 4))
    doc.build(story)
    return bio.getvalue()
